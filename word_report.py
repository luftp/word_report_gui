import docx
import pandas as pd
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docxcompose.composer import Composer
from docx import Document as Document_compose


class Word_report:
    """
    description:
                Classe escrita para auxiliar a elaboração do memorial de cálculo no word.
    """

    def __init__(self, path_template, from_template=True):

        if from_template:
            self.document = docx.Document(path_template)
            self.path = path_template

        self.niveis = self.id_carr = {1: 'EPC_Titulo1',
                                      2: 'EPC_Titulo2',
                                      3: 'EPC_Titulo3',
                                      4: 'EPC_Titulo4',
                                      5: 'EPC_Titulo5',
                                      6: 'EPC_Titulo6'}

    def add_item(self, texto, nivel=1):
        """
            description:
                Adiciona um item no memorial.
            Parameters:
                texto: texto a ser adicionado como item no arquivo.
                nivel: Nível de identação que o texto será inserido
        """
        self.document.add_paragraph(texto, style=self.niveis[nivel])

    def add_paragrafo(self, texto, style='EPC_ParNormal', breack_page=False):
        """
            description:
                Adiciona novo parágrafo no texto.
            Parameters:
                texto: Texto a ser inserido no documento.
                style: Estilo a ser usado no parâgrafo. Por padrão será usado o estilo EPC_ParNormal
        """
        self.document.add_paragraph(texto, style=style)

        if breack_page:
            self.document.add_page_break()

    @staticmethod
    def add_figure_caption(paragraph):
        """
            description:
                        Adiciona legenda a uma figura.
            Parameters:
                        paragraph: Onde será inserido a legenda.
        """
        run = paragraph.add_run()
        r = run._r
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'begin')
        r.append(fld_char)
        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'SEQ Figura \* ARABIC'
        r.append(instr_text)
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'end')
        r.append(fld_char)

    @staticmethod
    def add_table_caption(paragraph):
        """
            description:
                        Adiciona legenda a uma tabela.
            Parameters:
                        paragraph: Onde será inserido a legenda.
        """
        run = paragraph.add_run()
        r = run._r
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'begin')
        r.append(fld_char)
        instr_text = OxmlElement('w:instrText')
        instr_text.text = 'SEQ Tabela \* ARABIC'
        r.append(instr_text)
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'end')
        r.append(fld_char)

    def add_tabela(self, tabela: pd.DataFrame, legenda: str, style_legenda='EPC_Legenda',
                   style_tabela='EPC_Tabela_Style', breack_page=False):
        """
            description:
                        Adiciona uma tabela ao arquivo.
            Parameters:
                        tabela: Dataframe contendo a tabela a ser inserida na memória.
                        legenda: legenda a ser inserida na memória.
                        style_legenda: Estilo de texto a ser aplicada na legenda.
                        style_tabela: Estilo de tabela a ser usada na tabela inserida.
                        :param breack_page:
        """

        paragraph = self.document.add_paragraph('Tabela ', style=style_legenda)
        self.add_table_caption(paragraph)
        paragraph.add_run(' - ' + legenda)

        t = self.document.add_table(tabela.shape[0] + 1, tabela.shape[1])

        for j in range(tabela.shape[-1]):
            t.cell(0, j).text = tabela.columns[j]

        for i in range(tabela.shape[0]):
            for j in range(tabela.shape[-1]):
                t.cell(i + 1, j).text = str(tabela.values[i, j])

        t.style = style_tabela

        #self.document.add_paragraph('')
        par = self.document.add_paragraph('')
        run = par.add_run()
        if breack_page:
            run.add_break(WD_BREAK.PAGE)

    def add_figura(self, path_f, legenda, width=17, style_figura='EPC_Figura', style_legenda='EPC_Legenda',
                   breack_page=False, bookmark = None):
        """
            description:
                        Adiciona uma figura no arquivo.
            Parameters:
                        path_f: Endereço da figura acrescido de seu nome e extensão.
                        legenda: legenda a ser inserida sob a figura.
                        style_legenda: Estilo de texto a ser aplicado na legenda.
                        style_figura: Estilo de tabela a ser usada na tabela inserida.
                        breack_page: Boleano que indica se haverá quebra de página após a figura.
        """
        self.document.add_picture(path_f, width=Inches(width/2.54))
        # self.document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.document.paragraphs[-1].style = style_figura
        paragraph = self.document.add_paragraph('Figura ', style=style_legenda)

        self.add_figure_caption(paragraph)
        paragraph.add_run(' - ' + legenda)
        run = paragraph.add_run()

        if breack_page:
            run.add_break(WD_BREAK.PAGE)


        # MÉTODO ADICIONADO PARA FAZER REFERÊNCIA CRUZADA

        if bookmark != None:

            self.add_bookmark(paragraph, bookmark)

    # MÉTODO ADICIONADO PARA FAZER REFERÊNCIA CRUZADA

    def add_bookmark(self, paragraph, bookmark_name):
        """
        Adds a bookmark to a paragraph for cross-referencing.
        :param paragraph: The paragraph to bookmark.
        :param bookmark_name: The name of the bookmark.
        """
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), '0')
        bookmark_start.set(qn('w:name'), bookmark_name)

        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '0')

        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)



    def add_ref_to_figure(self, paragraph, bookmark_name):
        """
        Adds a REF field to reference the figure's SEQ field and display only the figure number.
        :param paragraph: The paragraph where the REF field will be added.
        :param bookmark_name: The bookmark name to reference.
        """
        paragraph = self.document.add_paragraph()

        run = paragraph.add_run()

        # Create the REF field to point to the SEQ field of the figure and show only the number
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = f' REF {bookmark_name} \h \# "0" '  # Use \# 0 to display only the number

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        # Append REF field parts to the paragraph
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


    @staticmethod
    def combine_all_docx(filename_master, files_list, path):
        """
            description:
                        Agrega arquivos a um arquivo master.
            Parameters:
                        filename_master: Endereço do arquivo master.
                        files_list: Lista de endereços dos arquivos a serem agregado.
                        path: Caminho do arquivo resultante.
        """
        number_of_sections = len(files_list)
        master = Document_compose(filename_master)
        composer = Composer(master)
        for i in range(0, number_of_sections):
            doc_temp = Document_compose(files_list[i])
            composer.append(doc_temp)
        composer.save(path)

    def save_file(self, path='-'):
        """
            description:
                Método por responsável por o arquivo.
            Parameters:
                path: Endereço para salvamento do modelo. Se nenhum endereço for pasasdo as modificações
                feitas no template serão salvas.
        """
        if path == '-':
            self.document.save(self.path)
        else:
            self.document.save(path)
